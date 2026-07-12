"""Shared helpers for building UPI-style survey research reports."""
import numpy as np
from scipy import stats
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SCALE = {"strongly disagree": 1, "disagree": 2, "neutral": 3, "agree": 4, "strongly agree": 5}

# ---------------- data / stats ----------------
def load_rows(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    return rows[0], rows[1:]

def to_score(v):
    if v is None:
        return np.nan
    return SCALE.get(str(v).strip().lower(), np.nan)

def likert_matrix(data, idxs):
    return np.array([[to_score(r[i]) for i in idxs] for r in data], dtype=float)

def item_stats(mat, questions):
    CATS = [("Strongly Agree", 5), ("Agree", 4), ("Neutral", 3), ("Disagree", 2), ("Strongly Disagree", 1)]
    out = []
    for j, q in enumerate(questions):
        col = mat[:, j]
        col = col[~np.isnan(col)]
        dist = {name: round(100.0 * np.sum(col == val) / len(col), 1) for name, val in CATS}
        out.append({"idx": j + 1, "q": q, "mean": round(float(col.mean()), 2),
                    "sd": round(float(col.std(ddof=1)), 2),
                    "agree_pct": round(100.0 * np.sum(col >= 4) / len(col), 1),
                    "dist": dist, "n": len(col)})
    return out

def cronbach_alpha(mat):
    m = mat[~np.isnan(mat).any(axis=1)]
    k = m.shape[1]
    if k < 2:
        return float("nan")
    item_var = m.var(axis=0, ddof=1).sum()
    total_var = m.sum(axis=1).var(ddof=1)
    if total_var == 0:
        return float("nan")
    return (k / (k - 1)) * (1 - item_var / total_var)

def cat_percent(data, idx):
    """Return ordered list of (value, count, pct) for a categorical column."""
    from collections import Counter
    vals = [str(r[idx]).strip() for r in data if r[idx] is not None]
    c = Counter(vals)
    total = len(vals)
    items = sorted(c.items(), key=lambda kv: -kv[1])
    return [(k, v, round(100.0 * v / total, 1)) for k, v in items], total

# ---------------- charts ----------------
def setup_style():
    plt.rcParams.update({"font.size": 10, "axes.grid": True, "grid.alpha": 0.3, "figure.dpi": 130})

BLUE = "#2f6fb0"
PALETTE = ["#1a5276", "#2e86c1", "#5dade2", "#85c1e9", "#aed6f1", "#d4e6f1"]

def bar_chart(path, labels, values, title, ylabel, ymax=None, color=BLUE, rotate=0, fmt="{:.2f}"):
    fig, ax = plt.subplots(figsize=(8, 4.5))
    bars = ax.bar(labels, values, color=color)
    ax.set_ylabel(ylabel); ax.set_title(title)
    if ymax: ax.set_ylim(0, ymax)
    if rotate: plt.xticks(rotation=rotate, ha="right")
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v + (ymax or max(values)) * 0.01, fmt.format(v),
                ha="center", va="bottom", fontsize=8)
    plt.tight_layout(); plt.savefig(path); plt.close()

def pie_chart(path, labels, values, title):
    fig, ax = plt.subplots(figsize=(6.5, 5))
    ax.pie(values, labels=labels, autopct="%1.1f%%", startangle=90,
           colors=PALETTE[:len(values)], textprops={"fontsize": 9})
    ax.set_title(title); ax.axis("equal")
    plt.tight_layout(); plt.savefig(path); plt.close()

def dim_bar(path, dim_names, means, title, neutral=3.0, highlight_idx=None):
    fig, ax = plt.subplots(figsize=(8, 4.5))
    cols = [BLUE] * len(means)
    if highlight_idx is not None:
        for i in highlight_idx: cols[i] = "#e67e22"
    bars = ax.bar(dim_names, means, color=cols)
    ax.axhline(neutral, color="grey", ls="--", lw=1, label=f"Neutral ({neutral})")
    ax.set_ylabel("Mean Score (1-5)"); ax.set_ylim(0, 5); ax.set_title(title)
    for b, v in zip(bars, means):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.05, f"{v:.2f}", ha="center", va="bottom")
    ax.legend(); plt.tight_layout(); plt.savefig(path); plt.close()

def group_bar(path, names, means, ns, title, ylabel="Mean Score (1-5)"):
    fig, ax = plt.subplots(figsize=(7.5, 4.5))
    bars = ax.bar(names, means, color=PALETTE[:len(names)])
    ax.set_ylabel(ylabel); ax.set_ylim(0, 5); ax.set_title(title)
    for b, v, nn in zip(bars, means, ns):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.05, f"{v:.2f}\n(n={nn})",
                ha="center", va="bottom", fontsize=9)
    plt.tight_layout(); plt.savefig(path); plt.close()

# ---------------- docx ----------------
class Report:
    def __init__(self):
        self.doc = Document()
        st = self.doc.styles["Normal"]
        st.font.name = "Calibri"; st.font.size = Pt(11)

    def p(self, text="", align=None, bold=False, size=None, italic=False):
        par = self.doc.add_paragraph(); run = par.add_run(text)
        run.bold = bold; run.italic = italic
        if size: run.font.size = Pt(size)
        if align == "c": par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "j": par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return par

    def h1(self, t): return self.doc.add_heading(t, level=1)
    def h2(self, t): return self.doc.add_heading(t, level=2)

    def para(self, t):
        par = self.doc.add_paragraph(t); par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; return par

    def bullet(self, t): self.doc.add_paragraph(t, style="List Bullet")
    def numbered(self, t): self.doc.add_paragraph(t, style="List Number")
    def pb(self): self.doc.add_page_break()

    def table(self, headers, rows, caption=None):
        if caption:
            cp = self.doc.add_paragraph(); cp.add_run(caption).bold = True
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, htxt in enumerate(headers):
            t.rows[0].cells[i].text = ""
            t.rows[0].cells[i].paragraphs[0].add_run(htxt).bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        return t

    def figure(self, path, caption):
        self.doc.add_picture(path, width=Inches(6.0))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp = self.doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption); r.italic = True; r.font.size = Pt(10)

    def title_page(self, title, degree, university, institution, month, name, regno):
        self.p(title, align="c", bold=True, size=16); self.p()
        self.p("A PROJECT REPORT", align="c", bold=True, size=12)
        self.p("Submitted in partial fulfilment of the requirements", align="c")
        self.p("for the award of the degree of", align="c")
        self.p(degree, align="c", bold=True, size=12)
        self.p(f"Affiliated to {university}", align="c")
        self.p(institution, align="c", bold=True)
        self.p(month, align="c", bold=True); self.p()
        self.p("Submitted by", align="c")
        self.p(name, align="c", bold=True, size=12)
        self.p(f"Register No : {regno}", align="c")
        self.pb()

    def save(self, path):
        self.doc.save(path)
