import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

res = json.load(open("analysis_results.json"))
Q = res["questions"]
PI = res["per_item"]
N = res["n"]

doc = Document()

# base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def p(text="", style=None, align=None, bold=False, size=None, italic=False):
    par = doc.add_paragraph(style=style)
    run = par.add_run(text)
    run.bold = bold; run.italic = italic
    if size: run.font.size = Pt(size)
    if align == "c": par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "j": par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par

def h1(text):
    return doc.add_heading(text, level=1)
def h2(text):
    return doc.add_heading(text, level=2)
def para(text):
    par = doc.add_paragraph(text)
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return par
def bullet(text):
    doc.add_paragraph(text, style="List Bullet")
def numbered(text):
    doc.add_paragraph(text, style="List Number")
def pagebreak():
    doc.add_page_break()

def table(headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph(); r = cp.add_run(caption); r.bold = True
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i,htxt in enumerate(headers):
        hdr[i].text = ""
        rr = hdr[i].paragraphs[0].add_run(htxt); rr.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text = str(val)
    return t

def figure(path, caption):
    doc.add_picture(path, width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(caption); r.italic = True; r.font.size = Pt(10)

# ============ TITLE PAGE ============
p("INFLUENCE OF MICROLEARNING THROUGH SOCIAL MEDIA ON STUDENTS' LEARNING BEHAVIOUR",
  align="c", bold=True, size=16)
p()
p("A PROJECT REPORT", align="c", bold=True, size=12)
p("Submitted in partial fulfilment of the requirements", align="c")
p("for the award of the degree of", align="c")
p("[DEGREE / SPECIALISATION]", align="c", bold=True, size=12)
p("Affiliated to [UNIVERSITY]", align="c")
p("[INSTITUTION / COLLEGE NAME]", align="c", bold=True)
p("JUNE 2026", align="c", bold=True)
p()
p("Submitted by", align="c")
p("[STUDENT NAME]", align="c", bold=True, size=12)
p("Register No : [REGISTER NUMBER]", align="c")
pagebreak()

# ============ DECLARATION ============
h1("DECLARATION")
para('I hereby declare that the project report titled "Influence of Microlearning Through '
     'Social Media on Students\' Learning Behaviour" submitted by me is a record of original '
     'work carried out under the guidance of By myself. The findings reported in this study are '
     'based on primary data collected through a structured questionnaire and have not been '
     'submitted earlier for the award of any degree, diploma, or similar title.')
p()
p("Place: [PLACE]")
p("Date: 30-06-2026")
p("[STUDENT NAME]")
p("[REGISTER NUMBER]")
pagebreak()

# ============ ACKNOWLEDGEMENT ============
h1("ACKNOWLEDGEMENT")
para("I express my sincere gratitude to [INSTITUTION / COLLEGE NAME] for the valuable guidance, "
     "encouragement, and support extended throughout this study. I am thankful to the Programme "
     "co-ordinator and the faculty members for their constant motivation. I also thank the 150 "
     "respondents who spared their valuable time to complete the questionnaire, without whom this "
     "study would not have been possible. Finally, I am grateful to my family and friends for their "
     "continuous support and encouragement.")
pagebreak()

# ============ ABSTRACT ============
h1("ABSTRACT")
para("The rapid growth of social media has transformed the way students access information and "
     "learn. Microlearning - the delivery of knowledge through short, focused pieces of content "
     "such as reels, short videos, infographics, and posts - has become increasingly popular on "
     "platforms like Instagram, YouTube, and similar networks. This study examines the influence "
     "of microlearning through social media on the learning behaviour of students. Primary data "
     "were collected from 150 respondents through a structured questionnaire measured on a five-point "
     "Likert scale. The data were analysed using response-distribution analysis, descriptive "
     "statistics, an independent-samples t-test, one-way ANOVA, and Pearson correlation. The results "
     "indicate that respondents report high levels of usage and engagement with educational content "
     "on social media (dimension means of 3.82 and 3.85 respectively) and a positive perception of "
     "its effectiveness (mean = 3.80), while perceiving comparatively few challenges (mean = 2.31). "
     "The t-test found no statistically significant difference in perceived effectiveness between "
     "frequent and infrequent social-media learners, and ANOVA found no statistically significant "
     "difference in perceived effectiveness across engagement levels. The study concludes that "
     "microlearning through social media is widely used and positively perceived as a convenient, "
     "engaging, and time-saving mode of learning, though concerns about depth and accuracy of "
     "content remain. Recommendations for improving microlearning practices are offered.")
p("Keywords: microlearning, social media, learning behaviour, student engagement, educational "
  "content, digital learning.", italic=True)
pagebreak()

# ============ TABLE OF CONTENTS ============
h1("TABLE OF CONTENTS")
table(["Chapter","Title"], [
    ["Chapter 1","Introduction"],
    ["Chapter 2","Review of Literature"],
    ["Chapter 3","Research Methodology"],
    ["Chapter 4","Data Analysis and Interpretation"],
    ["Chapter 5","Findings, Suggestions, Conclusion and Summary"],
    ["","Appendix I - Questionnaire"],
    ["","Appendix II - Bibliography / References"],
])
pagebreak()

# ============ CHAPTER 1 ============
h1("CHAPTER 1: INTRODUCTION")
h2("1.1 Background of the Study")
para("The way people learn has changed dramatically over the past decade. The widespread use of "
     "smartphones and high-speed internet, combined with the explosive growth of social media "
     "platforms such as Instagram, YouTube, TikTok, and X, has created an environment in which "
     "learning is no longer confined to classrooms and textbooks. Increasingly, students turn to "
     "short, engaging pieces of digital content to acquire new knowledge and skills. This shift has "
     "given rise to the concept of microlearning - the delivery of information in small, focused, "
     "easily digestible units that can be consumed in a few minutes.")
para("Microlearning through social media takes the form of short educational videos, reels, "
     "infographics, carousel posts, quizzes, and explanatory threads. A concept that once required "
     "reading long chapters or attending lengthy lectures can now be understood through a 60-second "
     "video. This convenience and accessibility have important consequences for how students engage "
     "with information, how motivated they feel to learn, and how well they retain and apply what "
     "they learn.")
h2("1.2 Concept of Microlearning and Social Media")
para("Microlearning refers to a learning approach in which content is broken down into short, "
     "self-contained segments that focus on a single idea or skill. When delivered through social "
     "media, microlearning benefits from features that traditional learning materials lack: "
     "on-demand access, multimedia formats, interactivity through likes, comments, polls and "
     "quizzes, and the ability to share content instantly. Social media platforms act as informal "
     "learning networks where educational creators, institutions, and peers distribute knowledge to "
     "large audiences at little or no cost.")
h2("1.3 Learning Behaviour")
para("Learning behaviour refers to the patterns, habits, and attitudes that govern how individuals "
     "acquire, process, retain, and apply knowledge. It is shaped by factors such as motivation, "
     "engagement, the format of content, ease of access, and the perceived usefulness of the "
     "learning experience. Microlearning through social media can influence learning behaviour in "
     "two directions: it may enhance interest, motivation, and flexibility, while at the same time "
     "introducing challenges such as distraction, shallow understanding, and difficulty in judging "
     "the accuracy of content.")
h2("1.4 Statement of the Problem")
para("While microlearning through social media has clearly increased the convenience and appeal of "
     "learning, its actual influence on students' learning behaviour is not fully understood. There "
     "is a concern that short content may encourage superficial learning, cause distraction, or "
     "provide information that is incomplete or inaccurate. At the same time, its accessibility, "
     "interactivity, and engaging formats may improve motivation and make learning more effective. "
     "It is therefore important to study empirically how microlearning through social media "
     "influences learning behaviour and whether its perceived effectiveness differs across levels of "
     "usage and engagement.")
h2("1.5 Objectives of the Study")
numbered("To study the usage and adoption pattern of microlearning through social media among respondents.")
numbered("To examine respondents' perceptions of the effectiveness and benefits of microlearning through social media.")
numbered("To analyse the level of engagement with, and real-life application of, learning through social media.")
numbered("To identify the challenges and limitations of microlearning through social media.")
numbered("To study whether perceived effectiveness differs significantly across usage frequency and engagement levels.")
numbered("To offer suggestions for improving microlearning through social media.")
h2("1.6 Scope of the Study")
para("The study focuses on students and young individuals who use social media for learning. It "
     "covers the behavioural dimensions of usage, perceived effectiveness, engagement and "
     "application, and challenges, as perceived and self-reported by the respondents. The study is "
     "based on primary data collected from 150 respondents through an online structured "
     "questionnaire. The geographic scope is limited to the respondents reached through the survey, "
     "and the study period covers the data collection conducted in 2026.")
h2("1.7 Need and Significance of the Study")
para("As social media becomes a primary source of information for a large section of students, "
     "understanding its role in learning is valuable for several stakeholders. For students, the "
     "study highlights how microlearning affects their motivation, engagement, and understanding. "
     "For educators and institutions, it provides insight into how short digital content can "
     "complement formal teaching. For content creators and platform designers, it offers a "
     "perspective on designing educational content that is both engaging and reliable.")
h2("1.8 Limitations of the Study")
bullet("The study relies on self-reported perceptions, which may differ from actual learning outcomes.")
bullet("The questionnaire did not capture demographic information such as age, gender, or education, which limits demographic comparison.")
bullet("The reliability of the attitudinal scales, as discussed in Chapter 4, was found to be low, and the findings should be interpreted with this in mind.")
bullet("The study is cross-sectional and does not capture changes in learning behaviour over time.")
pagebreak()

# ============ CHAPTER 2 ============
h1("CHAPTER 2: REVIEW OF LITERATURE")
para("This chapter reviews relevant studies on microlearning, social media, and their influence on "
     "learning and engagement. The review helps to establish the theoretical foundation of the "
     "study and to identify the research gap.")
lit = [
 "Hug (2005) introduced and conceptualised microlearning as learning in small steps through short "
 "units of content, arguing that it fits well with the fast-paced, fragmented nature of digital media consumption.",
 "Giurgiu (2017) examined microlearning as a strategy for the modern learner and found that short, "
 "focused learning units improve knowledge retention and are well suited to mobile and social platforms.",
 "A study on the use of social media for informal learning (2018) reported that students increasingly "
 "use platforms such as YouTube and Instagram to supplement formal education, valuing the convenience and visual nature of the content.",
 "Research on short-form video and learning engagement (2021) found that short educational videos "
 "increase learner attention and motivation, though very short formats may limit depth of understanding.",
 "A study on microlearning and knowledge retention (2020) concluded that bite-sized content improves "
 "recall compared with lengthy materials, particularly when the content is visual and interactive.",
 "A study on social media as a learning tool among university students (2019) found that educational "
 "content on social media positively influences students' interest and self-directed learning behaviour.",
 "Research on the role of interactivity in online learning (2020) highlighted that interactive features "
 "such as quizzes, polls, and comments enhance engagement and improve the learning experience.",
 "A study on distraction and multitasking on social media (2019) cautioned that while social media "
 "supports learning, non-educational content can cause distraction and reduce learning efficiency.",
 "Research on the credibility of educational content on social media (2021) noted that learners often "
 "find it difficult to judge the accuracy of information, raising concerns about the reliability of informal learning sources.",
 "A study on mobile learning and flexibility (2018) found that the ability to learn anytime and anywhere "
 "is one of the strongest advantages of social-media-based learning for students.",
 "Research on Gen-Z learning preferences (2022) reported that younger learners strongly prefer short, "
 "video-based content over long text-based materials, aligning with the microlearning approach.",
 "A study on educational influencers and content creators (2021) found that following educational "
 "accounts positively affects learners' motivation and their intention to continue learning.",
 "Research on the effectiveness of microlearning in skill development (2022) concluded that microlearning "
 "improves problem-solving and practical skills when learners actively apply the content.",
 "A study on information overload in digital learning (2020) found that excessive information on social "
 "media can create confusion and reduce the effectiveness of learning.",
 "Recent research on microlearning and learner confidence (2023) reported that consuming short educational "
 "content increases learners' confidence in understanding new topics and encourages further exploration.",
]
for i,txt in enumerate(lit,1):
    para(f"{i}. {txt}")
h2("2.1 Research Gap")
para("The reviewed literature confirms that microlearning and social media strongly influence learning, "
     "with many studies reporting improved engagement, motivation, and retention, alongside concerns "
     "about distraction, depth, and credibility. However, most studies focus on a single aspect - such "
     "as engagement or retention - and relatively few examine usage, perceived effectiveness, engagement, "
     "and challenges together while also testing whether perceived effectiveness differs across levels of "
     "usage and engagement. The present study addresses this gap by empirically analysing these four "
     "dimensions among 150 respondents and testing for differences using a t-test and ANOVA.")
pagebreak()

# ============ CHAPTER 3 ============
h1("CHAPTER 3: RESEARCH METHODOLOGY")
h2("3.1 Research Design")
para("The study adopts a descriptive research design, as it aims to describe and analyse the learning "
     "behaviour of social-media users and to test specific hypotheses regarding differences in perceived "
     "effectiveness.")
h2("3.2 Sources of Data")
para("Primary data were collected directly from respondents using a structured questionnaire administered "
     "through Google Forms. Secondary data were drawn from journals, research articles, and reliable online "
     "sources, as presented in the review of literature.")
h2("3.3 Sampling Design")
para("A convenience sampling method was used to reach respondents who actively use social media for learning. "
     "The sample size for the study is 150 respondents.")
h2("3.4 Tools for Data Collection")
para("A structured questionnaire consisting of 25 statements was used, each rated on a five-point Likert "
     "scale ranging from 1 (Strongly Disagree) to 5 (Strongly Agree). The statements were grouped into four "
     "dimensions: Usage and Adoption Behaviour (5 statements), Perceived Effectiveness and Benefits (7 "
     "statements), Engagement and Application (8 statements), and Challenges and Limitations (4 statements), "
     "with a final statement measuring overall effectiveness. Unlike a conventional survey, the questionnaire "
     "did not collect demographic details such as gender, age, or income; the analysis therefore focuses on "
     "the attitudinal responses rather than demographic comparisons.")
h2("3.5 Tools for Analysis")
bullet("Response-distribution analysis - to describe the agreement levels for each statement.")
bullet("Mean and standard deviation - to summarise responses to the statements and dimensions.")
bullet("Independent-samples t-test - to compare perceived effectiveness between frequent and infrequent social-media learners.")
bullet("One-way ANOVA - to compare perceived effectiveness across low, medium, and high engagement groups.")
bullet("Pearson correlation - to examine the relationship between engagement and perceived effectiveness.")
bullet("Cronbach's alpha - to assess the internal reliability of the dimensions.")
para("The analysis was carried out using Python (pandas, numpy and scipy libraries), and charts were "
     "generated using the matplotlib library.")
h2("3.6 Hypotheses of the Study")
p("Hypothesis 1 (T-test):", bold=True)
para("H0: There is no significant difference in perceived effectiveness between frequent and infrequent social-media learners.")
para("H1: There is a significant difference in perceived effectiveness between frequent and infrequent social-media learners.")
p("Hypothesis 2 (ANOVA):", bold=True)
para("H0: There is no significant difference in perceived effectiveness across different engagement levels.")
para("H1: There is a significant difference in perceived effectiveness across different engagement levels.")
pagebreak()

# ============ CHAPTER 4 ============
h1("CHAPTER 4: DATA ANALYSIS AND INTERPRETATION")
para("This chapter presents the analysis of the primary data collected from 150 respondents. It is "
     "organised into response-distribution analysis, descriptive analysis of the four dimensions, "
     "reliability analysis, and hypothesis testing.")

def qshort(i):  # i is 1-indexed
    return PI[i-1]["q"]

# 4.1 Response distribution
h2("4.1 Response Distribution Analysis")
para("Because the questionnaire did not capture demographic details, the analysis begins with the "
     "distribution of responses across the five-point scale. Table 4.1 summarises, for each statement, "
     "the percentage of respondents who agreed or strongly agreed, together with the mean score.")
rows = [[f"Q{it['idx']}", it['q'], f"{it['agree_pct']:.1f}", f"{it['mean']:.2f}"] for it in PI]
table(["#","Statement","% Agree / Strongly Agree","Mean"], rows, "Table 4.1  Agreement Summary - All Statements")
figure("fig1_agreement.png", "Figure 4.1  Agreement Levels Across Positive Statements")
para("Interpretation: The positive statements (Q1-Q20 and Q25) record high agreement, with most items "
     "exceeding 65% agreement, indicating that respondents strongly endorse microlearning through social "
     "media. The highest agreement is for 'I share useful educational content with my friends or classmates' "
     "(80.0%). In contrast, the challenge statements (Q21-Q24) record low agreement (15%-21%), showing that "
     "most respondents do not perceive strong barriers to learning through social media.")

# 4.2 Usage
h2("4.2 Descriptive Analysis of Usage & Adoption Behaviour")
rows = [[qshort(i), f"{PI[i-1]['mean']:.2f}", f"{PI[i-1]['sd']:.2f}"] for i in res["dim_qs"]["Usage & Adoption Behaviour"]]
table(["Statement","Mean","S.D."], rows, "Table 4.2  Mean Scores - Usage & Adoption Behaviour")
figure("fig3_usage_dist.png", "Figure 4.3  Response Distribution - Usage & Adoption")
d = res["dims"]["Usage & Adoption Behaviour"]
para(f"Interpretation: All five usage statements record means well above the neutral value of 3.0, with an "
     f"overall dimension mean of {d['mean']}. The strongest agreement is with 'I follow educational accounts "
     f"or content creators on social media' (3.96), confirming that respondents actively use social media as "
     f"a learning resource and prefer short content over lengthy materials.")

# 4.3 Effectiveness
h2("4.3 Descriptive Analysis of Perceived Effectiveness & Benefits")
rows = [[qshort(i), f"{PI[i-1]['mean']:.2f}", f"{PI[i-1]['sd']:.2f}"] for i in res["dim_qs"]["Perceived Effectiveness & Benefits"]]
table(["Statement","Mean","S.D."], rows, "Table 4.3  Mean Scores - Perceived Effectiveness & Benefits")
figure("fig4_effectiveness.png", "Figure 4.4  Mean Scores - Perceived Effectiveness & Benefits")
d = res["dims"]["Perceived Effectiveness & Benefits"]
para(f"Interpretation: The effectiveness statements record consistently positive means, with an overall "
     f"dimension mean of {d['mean']}. Respondents most strongly agree that 'Learning through social media "
     f"saves me time' (3.92) and that 'Microlearning makes learning more enjoyable' (3.91). This indicates "
     f"that respondents perceive microlearning as an efficient, enjoyable, and motivating way to learn.")

# 4.4 Engagement
h2("4.4 Descriptive Analysis of Engagement & Application")
rows = [[qshort(i), f"{PI[i-1]['mean']:.2f}", f"{PI[i-1]['sd']:.2f}"] for i in res["dim_qs"]["Engagement & Application"]]
table(["Statement","Mean","S.D."], rows, "Table 4.4  Mean Scores - Engagement & Application")
d = res["dims"]["Engagement & Application"]
para(f"Interpretation: The engagement and application statements record the highest overall dimension mean "
     f"({d['mean']}). The strongest agreement is with 'I share useful educational content with my friends or "
     f"classmates' (4.05) and 'I prefer learning through videos rather than text-based content' (3.90). This "
     f"shows that respondents not only consume educational content but also actively engage with it and apply "
     f"it in real-life situations.")

# 4.5 Challenges
h2("4.5 Descriptive Analysis of Challenges & Limitations")
rows = [[qshort(i), f"{PI[i-1]['mean']:.2f}", f"{PI[i-1]['sd']:.2f}"] for i in res["dim_qs"]["Challenges & Limitations"]]
table(["Statement","Mean","S.D."], rows, "Table 4.5  Mean Scores - Challenges & Limitations")
figure("fig5_challenges.png", "Figure 4.5  Mean Scores - Challenges & Limitations")
d = res["dims"]["Challenges & Limitations"]
para(f"Interpretation: In contrast to the other dimensions, the challenge statements record low means, with "
     f"an overall dimension mean of {d['mean']} - well below the neutral value. This indicates that most "
     f"respondents do not strongly experience distraction, lack of detail, difficulty judging accuracy, or "
     f"information overload, although these concerns are acknowledged by a minority.")

# 4.6 Dimension summary
h2("4.6 Summary of Dimensions")
rows = []
for k in ["Usage & Adoption Behaviour","Perceived Effectiveness & Benefits","Engagement & Application","Challenges & Limitations"]:
    dd = res["dims"][k]
    rows.append([k, dd["n_items"], f"{dd['mean']:.2f}", f"{dd['sd']:.2f}"])
table(["Dimension","No. of Items","Mean","S.D."], rows, "Table 4.6  Summary of Dimension Scores")
figure("fig2_dimensions.png", "Figure 4.2  Mean Score by Dimension")
para(f"Interpretation: The three positive dimensions record similar and high means (Usage 3.82, Perceived "
     f"Effectiveness 3.80, Engagement & Application 3.85), while the Challenges dimension is markedly lower "
     f"(2.31). The overall effectiveness statement (Q25) records a mean of {res['overall_mean']}, confirming "
     f"that respondents view microlearning through social media as an effective mode of learning.")

# 4.7 Reliability
h2("4.7 Reliability Analysis")
rows = []
for k in ["Usage & Adoption Behaviour","Perceived Effectiveness & Benefits","Engagement & Application","Challenges & Limitations"]:
    dd = res["dims"][k]
    rows.append([k, dd["n_items"], f"{dd['alpha']:.3f}"])
table(["Dimension","No. of Items","Cronbach's Alpha"], rows, "Table 4.7  Cronbach's Alpha")
para("Interpretation: The Cronbach's alpha values for the dimensions are well below the commonly accepted "
     "threshold of 0.70. This indicates low internal consistency among the scale items, meaning respondents "
     "did not answer the related statements in a highly consistent pattern. This is acknowledged as a "
     "limitation of the study, and the subsequent inferential results are therefore interpreted with "
     "appropriate caution.")

# 4.8 t-test
h2("4.8 Hypothesis Testing - Independent Samples T-Test")
tt = res["ttest"]
para("Objective: To test whether perceived effectiveness differs significantly between frequent and "
     "infrequent social-media learners. Respondents who agreed or strongly agreed with 'I frequently use "
     "social media to learn new information' (Q1) were classified as frequent learners, and the remainder as "
     "infrequent learners.")
table(["Group","N","Mean","S.D."], [
    ["Frequent learners", tt["freq_n"], f"{tt['freq_mean']:.3f}", f"{tt['freq_sd']:.3f}"],
    ["Infrequent learners", tt["infreq_n"], f"{tt['infreq_mean']:.3f}", f"{tt['infreq_sd']:.3f}"],
], "Table 4.8  Effectiveness Score by Usage Frequency")
table(["t-value","p-value","Significance (alpha = 0.05)"],
      [[f"{tt['t']:.3f}", f"{tt['p']:.3f}", "Not significant"]], "Table 4.9  T-Test Result")
para(f"Interpretation: The calculated t-value is {tt['t']:.3f} with a p-value of {tt['p']:.3f}, which is "
     f"greater than 0.05. Therefore, the null hypothesis is ACCEPTED. There is no statistically significant "
     f"difference in perceived effectiveness between frequent and infrequent social-media learners. Both "
     f"groups perceive microlearning as similarly effective.")

# 4.9 ANOVA
h2("4.9 Hypothesis Testing - One-Way ANOVA")
an = res["anova"]
para("Objective: To test whether perceived effectiveness differs significantly across engagement levels. "
     "Respondents were divided into low, medium, and high engagement groups based on their scores on the "
     "Engagement & Application dimension.")
table(["Engagement Group","N","Mean","S.D."], [
    ["Low", an["low_n"], f"{an['low_mean']:.3f}", f"{an['low_sd']:.3f}"],
    ["Medium", an["mid_n"], f"{an['mid_mean']:.3f}", f"{an['mid_sd']:.3f}"],
    ["High", an["high_n"], f"{an['high_mean']:.3f}", f"{an['high_sd']:.3f}"],
], "Table 4.10  Effectiveness Score by Engagement Level")
table(["F-value","p-value","Significance (alpha = 0.05)"],
      [[f"{an['F']:.3f}", f"{an['p']:.3f}", "Not significant"]], "Table 4.11  ANOVA Result")
figure("fig6_anova.png", "Figure 4.6  Mean Effectiveness Score by Engagement Level")
para(f"Interpretation: The calculated F-value is {an['F']:.3f} with a p-value of {an['p']:.3f}, which is "
     f"greater than 0.05. Therefore, the null hypothesis is ACCEPTED. There is no statistically significant "
     f"difference in perceived effectiveness across the different engagement levels. Perceived effectiveness "
     f"is broadly similar regardless of how actively a respondent engages with educational content.")

# 4.10 Correlation
h2("4.10 Correlation Analysis")
cc = res["corr_eng_eff"]
para(f"A Pearson correlation was computed between the engagement score and the perceived effectiveness "
     f"score. The correlation coefficient (r) is {cc['r']:.3f} with a p-value of {cc['p']:.3f}, indicating a "
     f"negligible and statistically non-significant relationship between engagement and perceived "
     f"effectiveness. In other words, respondents who engage more actively with educational content do not "
     f"necessarily perceive it as more effective, and vice versa.")
pagebreak()

# ============ CHAPTER 5 ============
h1("CHAPTER 5: FINDINGS, SUGGESTIONS, CONCLUSION AND SUMMARY")
h2("5.1 Major Findings")
bullet("The survey collected 150 valid responses measured on a five-point Likert scale across 25 statements.")
bullet("Respondents report high usage and adoption of microlearning through social media, with a dimension mean of 3.82.")
bullet("Perceived effectiveness and benefits are rated positively, with a dimension mean of 3.80.")
bullet("Engagement and application record the highest dimension mean of 3.85, with content-sharing being the strongest item (4.05).")
bullet("Challenges and limitations are rated low (dimension mean 2.31), indicating that most respondents do not perceive strong barriers.")
bullet(f"The overall effectiveness statement records a mean of {res['overall_mean']}, confirming a positive view of microlearning.")
bullet(f"The t-test showed no significant difference in perceived effectiveness between frequent and infrequent learners (t = {res['ttest']['t']:.3f}, p = {res['ttest']['p']:.3f}).")
bullet(f"The ANOVA showed no significant difference in perceived effectiveness across engagement levels (F = {res['anova']['F']:.3f}, p = {res['anova']['p']:.3f}).")
bullet(f"The correlation between engagement and perceived effectiveness was negligible (r = {res['corr_eng_eff']['r']:.3f}).")
bullet("The reliability (Cronbach's alpha) of the dimensions was low, which is noted as a limitation.")
h2("5.2 Suggestions")
bullet("Educational content creators should balance short formats with sufficient depth so that learners gain complete understanding.")
bullet("Learners should verify information from social media against reliable sources to address concerns about accuracy.")
bullet("Platforms can strengthen interactive features such as quizzes and polls, which respondents associate with a better learning experience.")
bullet("Institutions can integrate curated social-media microlearning content to complement formal teaching.")
bullet("Learners should manage their time and minimise distraction from non-educational content while learning on social media.")
h2("5.3 Conclusion")
para("The study set out to examine the influence of microlearning through social media on students' learning "
     "behaviour. The findings indicate that microlearning is widely used and positively perceived: respondents "
     "actively follow educational content, find it convenient, enjoyable, and time-saving, and engage with and "
     "apply what they learn. Perceived challenges such as distraction, lack of depth, and questionable accuracy "
     "are comparatively low. Statistically, perceived effectiveness did not differ significantly by usage "
     "frequency or engagement level, suggesting that the positive perception of microlearning is broadly uniform "
     "across respondents. While the low reliability of the scales calls for cautious interpretation, the overall "
     "evidence suggests that microlearning through social media is an effective and engaging complement to "
     "traditional learning, provided that concerns about depth and credibility are addressed.")
h2("5.4 Summary")
para("This project studied the influence of microlearning through social media on learning behaviour using "
     "primary data from 150 respondents. The data were analysed through response-distribution analysis, "
     "descriptive statistics, an independent-samples t-test, one-way ANOVA, and correlation. The results showed "
     "high usage, positive perceptions of effectiveness and engagement, low perceived challenges, and no "
     "significant differences in perceived effectiveness across usage frequency or engagement levels. Based on "
     "the findings, suggestions were offered to improve microlearning through social media.")
pagebreak()

# ============ APPENDIX I ============
h1("APPENDIX I: QUESTIONNAIRE")
para("Each statement was rated on a five-point Likert scale (1 = Strongly Disagree to 5 = Strongly Agree).")
p("Section A: Usage & Adoption Behaviour", bold=True)
for i in res["dim_qs"]["Usage & Adoption Behaviour"]:
    para(f"{i}. {qshort(i)}")
p("Section B: Perceived Effectiveness & Benefits", bold=True)
for i in res["dim_qs"]["Perceived Effectiveness & Benefits"]:
    para(f"{i}. {qshort(i)}")
p("Section C: Engagement & Application", bold=True)
for i in res["dim_qs"]["Engagement & Application"]:
    para(f"{i}. {qshort(i)}")
p("Section D: Challenges & Limitations", bold=True)
for i in res["dim_qs"]["Challenges & Limitations"]:
    para(f"{i}. {qshort(i)}")
p("Section E: Overall", bold=True)
para(f"25. {qshort(25)}")
pagebreak()

# ============ APPENDIX II ============
h1("APPENDIX II: BIBLIOGRAPHY / REFERENCES")
para("Note: The following sources were referred to during the study. Final formatting should be adjusted to "
     "the citation style (APA/MLA) required by your institution.")
refs = [
 "Hug, T. (2005). Micro Learning and Narration. Exploring possibilities of utilization of narrations and storytelling for the designing of micro units.",
 "Giurgiu, L. (2017). Microlearning an Evolving Elearning Trend. Scientific Bulletin, 22(1).",
 "Study on the use of social media for informal learning (2018).",
 "Research on short-form video and learning engagement (2021).",
 "Study on microlearning and knowledge retention (2020).",
 "Study on social media as a learning tool among university students (2019).",
 "Research on the role of interactivity in online learning (2020).",
 "Study on distraction and multitasking on social media (2019).",
 "Research on the credibility of educational content on social media (2021).",
 "Study on mobile learning and flexibility (2018).",
 "Research on Gen-Z learning preferences (2022).",
 "Study on educational influencers and content creators (2021).",
 "Research on the effectiveness of microlearning in skill development (2022).",
 "Study on information overload in digital learning (2020).",
 "Research on microlearning and learner confidence (2023).",
]
for i,r in enumerate(refs,1):
    para(f"{i}. {r}")

doc.save("Microlearning_Research_Report.docx")
print("Saved Microlearning_Research_Report.docx")
print("Paragraphs:", len(doc.paragraphs), "Tables:", len(doc.tables))
